import os
from datetime import date,timedelta
import docx
import pandas as pd
import logging
import sys

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('showcode_generator.log'),
        logging.StreamHandler(sys.stdout)
    ]
)

logger = logging.getLogger(__name__)

logger.info("Starting Show Code Generator v2")

df = pd.DataFrame(
    {"ISCI CODE (40x)" : ["SHOW CODE"],
     "CAMPAIGN (100x)" : ["X"],
     "TAG (100x)" : ["X"],
     "SPORT ID (40x)" : ["X"],
     "SUB CATEGORY (100x)" : ["X"],
     "YEAR (10x)" : ["X"],
     "LENGTH (5x)" : [10]})

fileName = "Encode Information "
bvsfileName = "BVS Encode Information "
titleFinal = "Encode Information week of "
TVFinal = ""
LiveFinal = ""
TVWGFinal = ""
LiveWGFinal = ""
startdate = date(2024, 8, 11)
startcode = 0
startlivecode = 0
tvWGstart = 0
liveWGstart = 0
FinalSheet = ""
tvCodes = []

# Create a document
doc = docx.Document()

# Add a paragraph to the document
p = doc.add_paragraph()

# Add some formatting to the paragraph
p.paragraph_format.line_spacing = 1.55
p.paragraph_format.space_after = 0

def addTitle(text):
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = docx.shared.Pt(11)

def addParagraph(text):
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = docx.shared.Pt(11)

def linebreak():
    run = p.add_run("\n")
    run.font.name = 'Calibri'
    run.font.size = docx.shared.Pt(11)

def convertDate(date):
    tempstring = date.strftime("%m") + "/" + date.strftime("%d") + "/" + date.strftime("%y") 
    return tempstring

def filenameDate(date):
    tempstring = date.strftime("%m") + date.strftime("%d") + date.strftime("%y")
    return tempstring

def outputDate(date):
    tempstring = date.strftime("%m") + date.strftime("%d") + date.strftime("%Y")
    return tempstring

def tvDailyGen(datestr, code):
    final = datestr + "\n:10 - s10a-tz" + str(code) + "-10\n"
    return final

def liveDailyGen(datestr, code):
    final = datestr + "\n:10 - s10a-tl" + str(code) + "-10\n"
    return final

def tvCodeGen(code):
    final = "s10a-tz" + code + "-10"
    return final

def liveCodeGen(code):
    final = "s10a-tl" + code + "-10"
    return final

def tvWGGen(datestr, dateend, snum, code):
    final = datestr + " - " + dateend + "\n:20 – s20g-tz" + str(snum) + "w" + str(code) + "–10\n:15 – s15g-tz" + str(snum) + "w" + str(code) + "-10\n:10 – s10g-tz" + str(snum) + "w" + str(code) + "-10\n:05 – s05g-tz" + str(snum) + "w" + str(code) + "-10"
    return final

def liveWGGen(datestr, dateend, snum, code):
    final = datestr + " - " + dateend + "\n:20 – s20g-tl" + str(snum) + "w" + str(code) + "–10\n:15 – s15g-tl" + str(snum) + "w" + str(code) + "-10\n:10 – s10g-tl" + str(snum) + "w" + str(code) + "-10\n:05 – s05g-tl" + str(snum) + "w" + str(code) + "-10"
    return final

def epNumberFix(num):
    if(len(str(num)) < 3):
        newnum = str(num)
        while(len(str(newnum)) < 3):
            newnum = "0"+str(newnum)
            #print(newnum)
        return newnum
    else:
        return num

def wgNumberFix(num):
    if(len(str(num)) < 2):
        newnum = str(num)
        while(len(str(newnum)) < 2):
            newnum = "0"+str(newnum)
            #print(newnum)
        return newnum
    else:
        return num

try:
    with open("setup.txt","r") as file:
        contents = file.read()
        logger.info(f"Successfully read setup.txt - Length: {len(contents)} characters")
        logger.debug(f"Setup file contents: {contents}")
except FileNotFoundError:
    logger.error("ERROR: setup.txt file not found in current directory")
    logger.error(f"Current working directory: {os.getcwd()}")
    logger.error("Please ensure setup.txt exists in the same directory as this script")
    sys.exit(1)
except PermissionError:
    logger.error("ERROR: Permission denied reading setup.txt")
    sys.exit(1)
except Exception as e:
    logger.error(f"ERROR: Unexpected error reading setup.txt: {e}")
    sys.exit(1)

#start
if(len(contents) == 181):
    logger.info("Setup file has correct length (181 characters). Processing...")
    
    try:
        #grab dates
        logger.info("Parsing dates from setup file...")
        year_str = contents[79:83]
        month_str = contents[75:77] 
        day_str = contents[77:79]
        logger.info(f"Date components - Year: '{year_str}', Month: '{month_str}', Day: '{day_str}'")
        
        startdate = date(int(year_str), int(month_str), int(day_str))
        logger.info(f"Successfully parsed start date: {startdate}")
        
    except ValueError as e:
        logger.error(f"ERROR: Invalid date in setup file - {e}")
        logger.error(f"Year: '{contents[79:83]}', Month: '{contents[75:77]}', Day: '{contents[77:79]}'")
        sys.exit(1)
    except Exception as e:
        logger.error(f"ERROR: Unexpected error parsing date - {e}")
        sys.exit(1)
        
    try:
        mon = convertDate(startdate)
        tempdate = startdate + timedelta(days=1)
        tues = convertDate(tempdate)
        tempdate = startdate + timedelta(days=2)
        wed = convertDate(tempdate)
        tempdate = startdate + timedelta(days=3)
        thurs = convertDate(tempdate)
        tempdate = startdate + timedelta(days=4)
        fri = convertDate(tempdate)
        tempdate = startdate + timedelta(days=4)
        fri = convertDate(tempdate)
        
        fileName += filenameDate(startdate) + " - " + filenameDate(tempdate) + ".docx"
        bvsfileName+= filenameDate(startdate) + " - " + filenameDate(tempdate) + ".xlsx"
        logger.info(f"Output filenames - Word: '{fileName}', Excel: '{bvsfileName}'")
        
        tempdate = startdate + timedelta(days=11)
        wgend = convertDate(tempdate)
        logger.info(f"Week dates - Mon: {mon}, Tue: {tues}, Wed: {wed}, Thu: {thurs}, Fri: {fri}")
        
    except Exception as e:
        logger.error(f"ERROR: Error generating dates - {e}")
        sys.exit(1)
    #print(wgend)

    try:
        #Create Title
        titleFinal += mon + " - " + fri
        logger.info(f"Title: '{titleFinal}'")

        #GrabTVCode
        tv_season = contents[95:97]
        tv_episode = contents[117:120]
        logger.info(f"TV Code components - Season: '{tv_season}', Episode: '{tv_episode}'")
        startcode = int(tv_season + tv_episode)
        logger.info(f"TV start code: {startcode}")

    except ValueError as e:
        logger.error(f"ERROR: Invalid TV code numbers in setup file - {e}")
        logger.error(f"TV Season (pos 95-97): '{contents[95:97]}', TV Episode (pos 117-120): '{contents[117:120]}'")
        sys.exit(1)
    except Exception as e:
        logger.error(f"ERROR: Error parsing TV codes - {e}")
        sys.exit(1)

    try:
        #CreateTVCodes
        logger.info("Generating TV codes...")
        TVFinal += tvDailyGen(mon,startcode) + tvDailyGen(tues,startcode+1) + tvDailyGen(wed,startcode+2) + tvDailyGen(thurs,startcode+3) + tvDailyGen(fri,startcode+4)
        tvCodes.append(tvCodeGen(str(startcode)))
        tvCodes.append(tvCodeGen(str(startcode+1)))
        tvCodes.append(tvCodeGen(str(startcode+2)))
        tvCodes.append(tvCodeGen(str(startcode+3)))
        tvCodes.append(tvCodeGen(str(startcode+4)))
        logger.info(f"Generated {len(tvCodes)} TV codes so far")

        #CreateLiveCodes
        logger.info("Generating Live codes...")
        live_season = contents[145:147]
        live_episode = contents[167:170]
        logger.info(f"Live Code components - Season: '{live_season}', Episode: '{live_episode}'")
        startlivecode = int(live_season + live_episode)
        logger.info(f"Live start code: {startlivecode}")
        
        LiveFinal += liveDailyGen(mon,startlivecode) + liveDailyGen(tues,startlivecode+1) + liveDailyGen(wed,startlivecode+2) + liveDailyGen(thurs,startlivecode+3) + liveDailyGen(fri,startlivecode+4)

    except ValueError as e:
        logger.error(f"ERROR: Invalid Live code numbers in setup file - {e}")
        logger.error(f"Live Season (pos 145-147): '{contents[145:147]}', Live Episode (pos 167-170): '{contents[167:170]}'")
        sys.exit(1)
    except Exception as e:
        logger.error(f"ERROR: Error generating TV/Live codes - {e}")
        sys.exit(1)
    #print(LiveFinal)

    #CreateTLWG
    liveWGstart = str(contents[179:181])
    liveseason = int(contents[145:147])
    LiveWGFinal += liveWGGen(fri,wgend,liveseason,liveWGstart)
    #print(LiveWGFinal)
    
    #CreateTzWG
    tvWGstart = str(contents[129:131])
    tzseason = int(contents[95:97])
    TVWGFinal += tvWGGen(fri,wgend,tzseason,tvWGstart)
    tvCodes.append("s20g-tz" + contents[95:97] + "w" + contents[129:131] + "–10")
    tvCodes.append("s15g-tz" + contents[95:97] + "w" + contents[129:131] + "–10")
    tvCodes.append("s10g-tz" + contents[95:97] + "w" + contents[129:131] + "–10")
    tvCodes.append("s05g-tz" + contents[95:97] + "w" + contents[129:131] + "–10")

    #append codes
    tvCodes.append(liveCodeGen(str(startlivecode)))
    tvCodes.append(liveCodeGen(str(startlivecode+1)))
    tvCodes.append(liveCodeGen(str(startlivecode+2)))
    tvCodes.append(liveCodeGen(str(startlivecode+3)))
    tvCodes.append(liveCodeGen(str(startlivecode+4)))
    
    tvCodes.append("s20g-tl" + contents[145:147] + "w" + contents[179:181] + "–10")
    tvCodes.append("s15g-tl" + contents[145:147] + "w" + contents[179:181] + "–10")
    tvCodes.append("s10g-tl" + contents[145:147] + "w" + contents[179:181] + "–10")
    tvCodes.append("s05g-tl" + contents[145:147] + "w" + contents[179:181] + "–10")
    
    #print(tvCodes)
    #print(TVWGFinal)
    #print(startdate.strftime("%A"))

    #Assembly
    addTitle(titleFinal)
    linebreak()
    linebreak()
    addTitle("TMZ 10 Promos:\n")
    addParagraph(TVFinal)
    linebreak()
    linebreak()
    addTitle("TMZ Live Promos:\n")
    addParagraph(LiveFinal)
    linebreak()
    linebreak()
    addTitle("TMZ Live Weekly Generic Promos:\n")
    addParagraph(LiveWGFinal)
    linebreak()
    linebreak()
    addTitle("TMZ Weekly Generic Promos:\n")
    addParagraph(TVWGFinal)

    try:
        logger.info("Saving Word document...")
        doc.save(str(fileName))
        logger.info(f"Successfully saved Word document: {fileName}")

        #Update next week
        logger.info("Updating setup file for next week...")
        WGNum = int(tvWGstart)+1
        TMZTVNUM = (int(contents[117:120])+6)
        TMZLIVENUM = (int(contents[167:170])+5)
        tempdate = startdate + timedelta(days=7)
        MondayDate = outputDate(tempdate)
        logger.info(f"Next week updates - WG: {WGNum}, TV Num: {TMZTVNUM}, Live Num: {TMZLIVENUM}, Date: {MondayDate}")
        
        contents = contents[:75] + MondayDate + contents[83:]
        contents = contents[:117] + str(epNumberFix(TMZTVNUM)) + contents[120:]
        contents = contents[:129] + str(wgNumberFix(WGNum)) + contents[131:]
        contents = contents[:179] + str(wgNumberFix(WGNum)) + contents[181:]
        contents = contents[:167] + str(epNumberFix(TMZLIVENUM)) + contents[170:]

        with open("setup.txt","w") as file:
            file.write(contents)
        logger.info("Successfully updated setup.txt for next week")

        for x in range(len(tvCodes)):
            df.loc[x+1] = [tvCodes[x], "X", "X", "X", "X", "X", int(tvCodes[x][1:3])]
        
        logger.info(f"Generated {len(tvCodes)} total show codes")
        logger.info("Script completed successfully!")
        
        #df.to_excel(bvsfileName, sheet_name='sheet1', index=False)
        
    except PermissionError as e:
        logger.error(f"ERROR: Permission denied saving files - {e}")
        logger.error("Check if files are open in another application or if you have write permissions")
        sys.exit(1)
    except Exception as e:
        logger.error(f"ERROR: Unexpected error during file operations - {e}")
        sys.exit(1)
    
else:
    logger.error(f"ERROR: Setup file has incorrect length!")
    logger.error(f"Expected: 181 characters, Got: {len(contents)} characters")
    logger.error("The setup.txt file must be exactly 181 characters long")
    logger.error(f"Current contents preview: '{contents[:50]}...'")
    sys.exit(1)


tempdate = startdate + timedelta(days=1)
