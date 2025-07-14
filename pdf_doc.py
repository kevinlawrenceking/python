import fitz
import sys
import os
from collections import defaultdict
from docx import Document
from docx.shared import Pt, Inches

def rect_overlap(rect1, rect2):
    return not (
        rect1.x1 < rect2.x0 or rect1.x0 > rect2.x1 or
        rect1.y1 < rect2.y0 or rect1.y0 > rect2.y1
    )

def extract_highlighted_text_layout_aware(doc):
    results = defaultdict(list)
    for page_number in range(len(doc)):
        page = doc[page_number]
        wordlist = page.get_text("words")
        highlight_rects = []
        for annot in page.annots():
            if annot.type[0] != 8:
                continue
            quads = [annot.vertices[i:i+4] for i in range(0, len(annot.vertices), 4)]
            for quad in quads:
                r = fitz.Quad(quad).rect
                highlight_rects.append(r)
        if not highlight_rects:
            continue
        matched_words = []
        for w in wordlist:
            word_rect = fitz.Rect(w[:4])
            if any(rect_overlap(word_rect, h) for h in highlight_rects):
                matched_words.append(w)
        matched_words.sort(key=lambda w: (round(w[1], 1), w[0]))
        lines = []
        current_line_y = None
        current_line = []
        for w in matched_words:
            y = round(w[1], 1)
            if current_line_y is None or abs(y - current_line_y) <= 2:
                current_line.append(w[4])
                current_line_y = y
            else:
                lines.append(" ".join(current_line))
                current_line = [w[4]]
                current_line_y = y
        if current_line:
            lines.append(" ".join(current_line))
        clean_text = "\n".join(line.strip() for line in lines if line.strip())
        if clean_text:
            results[page_number + 1].append(clean_text)
    return results

def write_to_docx(highlight_dict, output_path):
    doc = Document()
    doc.add_heading("Transcript Highlights", 0)

    # Set left/right margins to 1.5 inches
    for section in doc.sections:
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.5)

    for page, blocks in sorted(highlight_dict.items()):
        doc.add_heading(f"Page {page}", level=2)
        for block in blocks:
            para = doc.add_paragraph(block)
            # Set all runs in the paragraph to 10.5pt (~14px)
            for run in para.runs:
                run.font.size = Pt(14)
        doc.add_paragraph("")  # blank line after each page

    doc.save(output_path)

def main():
    if len(sys.argv) < 3:
        print("Usage: pdf.py input.pdf output.docx")
        return

    INPUT_PDF = sys.argv[1]
    OUTPUT_DOCX = sys.argv[2]

    if not os.path.isfile(INPUT_PDF):
        print(f"File not found: {INPUT_PDF}")
        return

    doc = fitz.open(INPUT_PDF)
    highlights = extract_highlighted_text_layout_aware(doc)
    doc.close()
    if not highlights:
        print("No highlighted text found.")
        return

    # Ensure output directory exists
    output_dir = os.path.dirname(OUTPUT_DOCX)
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    write_to_docx(highlights, OUTPUT_DOCX)
    print(f"Formatted highlights saved to {OUTPUT_DOCX}")

if __name__ == "__main__":
    main()
